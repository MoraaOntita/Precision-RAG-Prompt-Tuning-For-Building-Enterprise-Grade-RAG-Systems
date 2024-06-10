# tests/test_system.py
import os
import json
import tempfile
import pytest
from docx import Document
from streamlit import session_state as st_session_state
from scripts.prompt_generation.retriever import retrieve_relevant_context
from scripts.prompt_generation.prompt_generator import generate_prompts
from scripts.automate_evaluation_service.evaluator import evaluate_test_cases
from scripts.automate_evaluation_service.report_generator import generate_report

# Utility function to create a sample document
def create_sample_docx():
    doc = Document()
    doc.add_heading('Sample Document', 0)
    doc.add_paragraph('Business Need: This project aims to solve a critical business problem.')
    doc.add_paragraph('Background Context: The project is set in the context of a rapidly growing industry.')
    doc.add_paragraph('Learning Outcome: Participants will gain hands-on experience.')
    doc.add_paragraph('Deliverables: Weekly reports, final presentation.')
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

# Test function for the entire system
def test_system():
    # Step 1: Create a sample document
    sample_doc_path = create_sample_docx()

    # Step 2: Extract text from the document
    doc = Document(sample_doc_path)
    document_text = ""
    for paragraph in doc.paragraphs:
        document_text += paragraph.text + "\n"
    assert len(document_text) > 0, "Document extraction failed"

    # Step 3: Generate prompts based on document text
    prompt = "Generate a summary for the business need section."
    relevant_contexts = retrieve_relevant_context(prompt)
    best_prompt, ranked_prompts = generate_prompts(prompt, relevant_contexts, os.getenv('OPENAI_API_KEY'))
    assert best_prompt is not None, "Prompt generation failed"
    assert len(ranked_prompts) > 0, "No ranked prompts generated"

    # Step 4: Evaluate the generated prompts
    test_cases = [
        "What is the main goal of the business need section?",
        "Describe the learning outcomes expected from this challenge.",
        "What are the deliverables for this week?",
        "Summarize the background context of this challenge.",
        "List the instructions provided for the submissions."
    ]
    evaluation_results = evaluate_test_cases(test_cases, [best_prompt] + ranked_prompts)
    assert len(evaluation_results) > 0, "Evaluation of prompts failed"

    # Step 5: Generate a report from the evaluation results
    report = generate_report(evaluation_results)
    assert "Test Case Evaluation Report" in report, "Report generation failed"

    # Optional: Save the report to a temporary file for manual inspection
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as temp_report_file:
        temp_report_file.write(report.encode('utf-8'))
        temp_report_path = temp_report_file.name

    # Step 6: Simulate user interaction with the Streamlit UI
    st_session_state.prompt_data = []

    # Simulate user input
    st_session_state.prompt_data.append({
        "document": os.path.basename(sample_doc_path),
        "prompt": prompt,
        "expected_output": "A concise summary of the business need."
    })

    # Save prompt data to a temporary JSON file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".json") as temp_json_file:
        json.dump(st_session_state.prompt_data, temp_json_file, indent=4)
        temp_json_path = temp_json_file.name

    assert os.path.exists(temp_json_path), "Prompt data JSON file was not created"

    # Print paths for manual inspection if necessary
    print(f"Sample document path: {sample_doc_path}")
    print(f"Report path: {temp_report_path}")
    print(f"Prompt data JSON path: {temp_json_path}")

if __name__ == "__main__":
    test_system()
